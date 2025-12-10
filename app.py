import streamlit as st
import sqlite3
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO

# ==========================================
# 1. æ•°æ®åº“é…ç½®ä¸å·¥å…·å‡½æ•°
# ==========================================

DB_FILE = "fire_inspections.db"

def init_db():
    """åˆå§‹åŒ–æ•°æ®åº“è¡¨"""
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute('''
              CREATE TABLE IF NOT EXISTS inspections (
                                                         id INTEGER PRIMARY KEY AUTOINCREMENT,
                                                         project_name TEXT NOT NULL,
                                                         category TEXT,
                                                         loc TEXT,
                                                         desc TEXT,
                                                         remark TEXT,
                                                         img_bytes BLOB,
                                                         created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
              )
              ''')
    conn.commit()
    conn.close()

def get_all_projects():
    """è·å–æ‰€æœ‰å”¯ä¸€çš„é¡¹ç›®åç§°"""
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("SELECT DISTINCT project_name FROM inspections ORDER BY created_at DESC")
    projects = [row[0] for row in c.fetchall()]
    conn.close()
    if not projects:
        return ["é»˜è®¤é¡¹ç›®"]
    return projects

def add_item_to_db(project, category, loc, desc, remark, img_bytes):
    """æ·»åŠ ä¸€æ¡è®°å½•"""
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute('''
              INSERT INTO inspections (project_name, category, loc, desc, remark, img_bytes)
              VALUES (?, ?, ?, ?, ?, ?)
              ''', (project, category, loc, desc, remark, img_bytes))
    conn.commit()
    conn.close()

def get_items_by_project(project_name):
    """è·å–æŒ‡å®šé¡¹ç›®çš„æ‰€æœ‰è®°å½• (æŒ‰æ—¶é—´å€’åº)"""
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("SELECT * FROM inspections WHERE project_name = ? ORDER BY id DESC", (project_name,))
    rows = c.fetchall()
    conn.close()

    data_list = []
    for row in rows:
        data_list.append({
            "id": row["id"],
            "category": row["category"],
            "desc": row["desc"],
            "loc": row["loc"],
            "remark": row["remark"],
            "img_bytes": row["img_bytes"]
        })
    return data_list

def delete_item_from_db(item_id):
    """æ ¹æ®IDåˆ é™¤å•æ¡é—®é¢˜è®°å½•"""
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("DELETE FROM inspections WHERE id = ?", (item_id,))
    conn.commit()
    conn.close()

def delete_project_from_db(project_name):
    """[æ–°å¢] åˆ é™¤æ•´ä¸ªé¡¹ç›®çš„æ‰€æœ‰è®°å½•"""
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("DELETE FROM inspections WHERE project_name = ?", (project_name,))
    conn.commit()
    conn.close()

# å¯åŠ¨æ—¶ç¡®ä¿æ•°æ®åº“å­˜åœ¨
init_db()

# ==========================================
# 2. Word ç”Ÿæˆé€»è¾‘ (ä¿æŒä¸å˜)
# ==========================================
def set_font(run, font_name='å®‹ä½“', size=10, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.bold = bold

def create_word_file(report_name, data_list):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run(f'{report_name} - æ¶ˆé˜²æ£€æŸ¥é—®é¢˜æ¸…å•')
    set_font(run_title, 'é»‘ä½“', 18, bold=True)

    categories = ["å»ºç­‘é˜²ç«é—®é¢˜æ¸…å•", "æ¶ˆé˜²è®¾æ–½é—®é¢˜æ¸…å•"]

    for cat_name in categories:
        current_items = [item for item in data_list if item['category'] == cat_name]
        doc.add_paragraph("")
        prefix = "ä¸€ã€" if cat_name == "å»ºç­‘é˜²ç«é—®é¢˜æ¸…å•" else "äºŒã€"
        h_p = doc.add_paragraph()
        run_h = h_p.add_run(f"{prefix}{cat_name}")
        set_font(run_h, 'é»‘ä½“', 14, bold=True)

        if not current_items:
            p_none = doc.add_paragraph("ï¼ˆè¯¥é¡¹æ— é—®é¢˜ï¼‰")
            p_none.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.autofit = False
        widths = [Cm(1.5), Cm(7), Cm(6), Cm(2.5)]

        headers = ["åºå·", "é—®é¢˜æè¿°", "ç›¸å…³ç…§ç‰‡", "å¤‡æ³¨"]
        hdr_cells = table.rows[0].cells
        for i, text in enumerate(headers):
            hdr_cells[i].width = widths[i]
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_font(run, 'å®‹ä½“', 12, bold=True)

        for idx, item in enumerate(current_items, 1):
            row_cells = table.add_row().cells

            p1 = row_cells[0].paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p1.add_run(str(idx)))

            p2 = row_cells[1].paragraphs[0]
            run_desc = p2.add_run(f"é—®é¢˜æè¿°ï¼š{item['desc']}\n")
            set_font(run_desc)
            run_loc = p2.add_run(f"é—®é¢˜ä½ç½®ï¼š{item['loc']}")
            set_font(run_loc)

            cell_img = row_cells[2]
            p3 = cell_img.paragraphs[0]
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if item['img_bytes']:
                try:
                    run_img = p3.add_run()
                    run_img.add_picture(BytesIO(item['img_bytes']), width=Inches(2.0))
                except:
                    set_font(p3.add_run("[å›¾ç‰‡æ ¼å¼é”™è¯¯]"))
            else:
                set_font(p3.add_run("/"))

            p4 = row_cells[3].paragraphs[0]
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p4.add_run(item['remark']))

    return doc

# ==========================================
# 3. é¡µé¢ UI ä¸äº¤äº’é€»è¾‘
# ==========================================
st.set_page_config(page_title="æ¶ˆé˜²æ£€æŸ¥åŠ©æ‰‹", layout="centered")

# --- çŠ¶æ€ç®¡ç†ä¸æ•°æ®åŠ è½½é€»è¾‘ ---

if 'current_report_name' not in st.session_state:
    st.session_state.current_report_name = "é»˜è®¤é¡¹ç›®"

# è¯»å–æ•°æ®åº“é‡Œçš„é¡¹ç›®
db_projects = get_all_projects()

# ç¡®ä¿å½“å‰é¡¹ç›®åœ¨åˆ—è¡¨ä¸­
current_name = st.session_state.current_report_name
if current_name not in db_projects:
    db_projects.insert(0, current_name)

# --- é¡¶éƒ¨ï¼šé¡¹ç›®ç®¡ç†åŒºåŸŸ ---
with st.expander(f"ğŸ“‚ å½“å‰é¡¹ç›®ï¼š{current_name} (ç‚¹å‡»åˆ‡æ¢/åˆ é™¤)", expanded=False):

    # ä½¿ç”¨ Tabs åˆ†å‰²â€œåˆ‡æ¢/æ–°å»ºâ€å’Œâ€œåˆ é™¤â€åŠŸèƒ½ï¼Œç•Œé¢æ›´æ•´æ´
    tab_switch, tab_manage = st.tabs(["ğŸ” åˆ‡æ¢æˆ–æ–°å»º", "âš ï¸ åˆ é™¤é¡¹ç›®"])

    # --- Tab 1: åˆ‡æ¢ä¸æ–°å»º ---
    with tab_switch:
        # é€‰æ‹©æ¡†
        selected_report = st.selectbox(
            "é€‰æ‹©ç°æœ‰é¡¹ç›®",
            db_projects,
            index=db_projects.index(current_name)
        )

        if selected_report != current_name:
            st.session_state.current_report_name = selected_report
            st.rerun()

        st.write("---")
        # æ–°å»ºé¡¹ç›®é€»è¾‘
        col_new_1, col_new_2 = st.columns([3, 1])
        with col_new_1:
            new_report_name = st.text_input("æ–°å»ºé¡¹ç›®åç§°", placeholder="è¾“å…¥é¡¹ç›®å (å¦‚ï¼šä¸‡è¾¾å¹¿åœº)", label_visibility="collapsed")
        with col_new_2:
            if st.button("æ–°å»ºå¹¶åˆ‡æ¢"):
                if new_report_name and new_report_name.strip():
                    st.session_state.current_report_name = new_report_name.strip()
                    st.rerun()
                else:
                    st.warning("åç§°ä¸èƒ½ä¸ºç©º")

    # --- Tab 2: åˆ é™¤é¡¹ç›® (æ–°å¢åŠŸèƒ½) ---
    with tab_manage:
        st.write(f"æ­£åœ¨ç®¡ç†ï¼š**{current_name}**")

        if current_name == "é»˜è®¤é¡¹ç›®":
            st.info("ğŸš« ã€é»˜è®¤é¡¹ç›®ã€‘æ˜¯ç³»ç»Ÿä¿ç•™é¡¹ï¼Œä¸å¯åˆ é™¤ã€‚")
        else:
            # é˜²è¯¯åˆ æœºåˆ¶ï¼šå…ˆå‹¾é€‰ï¼Œå†æ˜¾ç¤ºåˆ é™¤æŒ‰é’®
            confirm_del = st.checkbox(f"æˆ‘ç¡®è®¤è¦åˆ é™¤é¡¹ç›®ã€{current_name}ã€‘åŠå…¶æ‰€æœ‰æ•°æ®", key="del_chk")

            if confirm_del:
                if st.button("ğŸ—‘ï¸ ç¡®è®¤å½»åº•åˆ é™¤", type="primary"):
                    # 1. ä»æ•°æ®åº“åˆ é™¤æ•°æ®
                    delete_project_from_db(current_name)

                    # 2. çŠ¶æ€é‡ç½®ï¼šåˆ é™¤åå›åˆ°â€œé»˜è®¤é¡¹ç›®â€æˆ–åˆ—è¡¨é‡Œçš„ç¬¬ä¸€ä¸ª
                    remaining = [p for p in db_projects if p != current_name]
                    fallback_project = remaining[0] if remaining else "é»˜è®¤é¡¹ç›®"

                    st.session_state.current_report_name = fallback_project

                    st.success(f"é¡¹ç›® {current_name} å·²åˆ é™¤ï¼")
                    st.rerun()

# --- åŠ è½½å½“å‰é¡¹ç›®æ•°æ® ---
current_list = get_items_by_project(current_name)

# --- æ ¸å¿ƒåŒºåŸŸï¼šæ·»åŠ é—®é¢˜ ---
st.markdown("### ğŸ“¸ ç°åœºå½•å…¥")

with st.container(border=True):
    with st.form("mobile_add_form", clear_on_submit=True):
        col1, col2 = st.columns([2, 1])
        with col1:
            location = st.text_input("ğŸ“ é—®é¢˜ä½ç½®", placeholder="å¦‚ï¼š8æ¥¼æ¥¼æ¢¯é—´")
        with col2:
            st.write("") # å ä½
            st.caption(f"å½“å‰: {current_name}")

        category = st.radio("âš ï¸ é—®é¢˜ç±»åˆ«", ["å»ºç­‘é˜²ç«é—®é¢˜æ¸…å•", "æ¶ˆé˜²è®¾æ–½é—®é¢˜æ¸…å•"], horizontal=True)
        desc = st.text_area("ğŸ“ é—®é¢˜æè¿°", placeholder="æè¿°å…·ä½“éšæ‚£...", height=100)

        st.markdown("**ğŸ“· æ·»åŠ ç…§ç‰‡ (ä»»é€‰ä¸€ç§)**")
        col_cam, col_upl = st.tabs(["è°ƒç”¨æ‘„åƒå¤´", "ä»ç›¸å†Œä¸Šä¼ "])
        with col_cam:
            camera_file = st.camera_input("ç‚¹å‡»æ‹ç…§", label_visibility="collapsed")
        with col_upl:
            uploaded_file = st.file_uploader("é€‰æ‹©æ–‡ä»¶", type=['png', 'jpg', 'jpeg'], label_visibility="collapsed")

        remark = st.text_input("ğŸ’¡ å¤‡æ³¨ (é€‰å¡«)", placeholder="æ•´æ”¹äºº/å»ºè®®")

        submitted = st.form_submit_button("âœ… ç¡®è®¤æ·»åŠ ", use_container_width=True, type="primary")

        if submitted:
            if not desc or not location:
                st.error("ä½ç½®å’Œæè¿°å¿…å¡«ï¼")
            else:
                final_img = camera_file if camera_file else uploaded_file
                img_data = final_img.getvalue() if final_img else None

                add_item_to_db(current_name, category, location, desc, remark, img_data if img_data else b'')

                st.success("å·²ä¿å­˜ï¼")
                st.rerun()

# --- åˆ—è¡¨å±•ç¤ºåŒº ---
st.markdown("---")
st.markdown(f"### ğŸ“‹ å·²è®°å½• ({len(current_list)})")

if not current_list:
    st.info(f"é¡¹ç›®ã€{current_name}ã€‘æš‚æ— è®°å½•ï¼Œè¯·åœ¨ä¸Šæ–¹æ·»åŠ ã€‚")
else:
    for item in current_list:
        with st.container(border=True):
            col_top_1, col_top_2 = st.columns([3, 1])
            with col_top_1:
                st.markdown(f"**ğŸ“ {item['loc']}**")
            with col_top_2:
                tag_color = "red" if "é˜²ç«" in item['category'] else "orange"
                st.caption(f":{tag_color}[{item['category'][:4]}]")

            st.text(item['desc'])

            if item['img_bytes']:
                st.image(item['img_bytes'], width=150)

            col_foot_1, col_foot_2 = st.columns([3, 1])
            with col_foot_1:
                if item['remark']:
                    st.caption(f"å¤‡æ³¨: {item['remark']}")
            with col_foot_2:
                if st.button("ğŸ—‘ï¸", key=f"del_{item['id']}"):
                    delete_item_from_db(item['id'])
                    st.rerun()

# --- åº•éƒ¨ï¼šä¸‹è½½åŒºåŸŸ ---
st.markdown("---")

if current_list:
    doc_object = create_word_file(current_name, current_list[::-1])
    output_buffer = BytesIO()
    doc_object.save(output_buffer)
    output_buffer.seek(0)

    st.download_button(
        label="ğŸ“¥ ç”Ÿæˆå¹¶ä¸‹è½½ Word æŠ¥å‘Š",
        data=output_buffer,
        file_name=f"{current_name}_æ¶ˆé˜²é—®é¢˜æ¸…å•.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
        type="primary"
    )
else:
    st.caption("æš‚æ— æ•°æ®å¯ä¸‹è½½")

st.write("")
st.write("")